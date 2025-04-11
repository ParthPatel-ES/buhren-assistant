"""
Document Processor for German language documents.

This module handles parsing, extraction, and chunking of text from
German documents in various formats (.docx, .pdf).
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass

# Document parsing libraries
import docx
import pypdf
from langchain.text_splitter import RecursiveCharacterTextSplitter

# German language processing
import spacy

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Structured container for document metadata."""
    source: str
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    pages: Optional[int] = None
    language: str = "de"  # Default to German
    doc_type: Optional[str] = None

@dataclass
class DocumentChunk:
    """Represents a chunk of text with its metadata."""
    text: str
    metadata: DocumentMetadata
    chunk_id: int


class DocumentProcessor:
    """
    Processes German documents (.docx, .pdf) for RAG applications.
    
    Handles document loading, text extraction, cleaning, and chunking
    with special consideration for German language characteristics.
    """
    
    def __init__(
        self, 
        chunk_size: int = 512, 
        chunk_overlap: int = 50,
        spacy_model: str = "de_core_news_sm"
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between consecutive chunks in characters
            spacy_model: German language model for NLP processing
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Load German language model for better text processing
        try:
            self.nlp = spacy.load(spacy_model)
            logger.info(f"Loaded German spaCy model: {spacy_model}")
        except Exception as e:
            logger.warning(f"Could not load spaCy model: {e}")
            logger.warning("Run 'python -m spacy download de_core_news_sm' to install")
            self.nlp = None
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of document chunks with metadata
        """
        all_chunks = []
        
        try:
            for root, _, files in os.walk(directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    if file_ext == '.pdf':
                        chunks = self.process_pdf(file_path)
                        all_chunks.extend(chunks)
                        logger.info(f"Processed PDF: {file_path}, extracted {len(chunks)} chunks")
                    
                    elif file_ext == '.docx':
                        chunks = self.process_docx(file_path)
                        all_chunks.extend(chunks)
                        logger.info(f"Processed DOCX: {file_path}, extracted {len(chunks)} chunks")
                    
                    else:
                        logger.debug(f"Skipping unsupported file type: {file_path}")
        
        except Exception as e:
            logger.error(f"Error processing directory {directory_path}: {str(e)}")
        
        return all_chunks
    
    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """
        Extract text and metadata from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of document chunks with metadata
        """
        try:
            with open(file_path, 'rb') as f:
                pdf = pypdf.PdfReader(f)
                
                # Extract metadata
                metadata = DocumentMetadata(
                    source=file_path,
                    title=pdf.metadata.get('/Title', os.path.basename(file_path)),
                    author=pdf.metadata.get('/Author'),
                    created_date=pdf.metadata.get('/CreationDate'),
                    pages=len(pdf.pages),
                    doc_type='pdf'
                )
                
                # Extract text from all pages
                text = ""
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"{page_text}\nPage: {page_num + 1}\n\n"
                
                # Process and clean the extracted text
                clean_text = self._clean_text(text)
                
                # Create chunks
                return self._create_chunks(clean_text, metadata)
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return []
    
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """
        Extract text and metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of document chunks with metadata
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata from document properties
            prop = doc.core_properties
            metadata = DocumentMetadata(
                source=file_path,
                title=prop.title or os.path.basename(file_path),
                author=prop.author,
                created_date=str(prop.created) if prop.created else None,
                doc_type='docx'
            )
            
            # Extract text while preserving some structure
            text = ""
            
            # Process headings and paragraphs
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    # Add extra newlines for headings to improve chunking
                    heading_level = para.style.name.replace('Heading', '').strip()
                    text += f"\n\n{'#' * int(heading_level) if heading_level.isdigit() else '#'} {para.text}\n\n"
                else:
                    text += para.text + "\n"
            
            # Process tables
            for table in doc.tables:
                text += "\nTable:\n"
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
                text += "\n"
            
            # Clean and process the text
            clean_text = self._clean_text(text)
            
            # Create chunks
            return self._create_chunks(clean_text, metadata)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return []
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize German text.
        
        Args:
            text: Raw text from document
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize German umlauts if they appear in decomposed form
        text = text.replace('a\u0308', 'ä')
        text = text.replace('o\u0308', 'ö')
        text = text.replace('u\u0308', 'ü')
        text = text.replace('A\u0308', 'Ä')
        text = text.replace('O\u0308', 'Ö')
        text = text.replace('U\u0308', 'Ü')
        
        # Fix common OCR errors in German texts
        text = text.replace('fi', 'fi')  # Fix common ligature issue
        text = text.replace('fl', 'fl')  # Fix common ligature issue
        
        # Use spaCy for additional German text normalization if available
        if self.nlp:
            doc = self.nlp(text)
            # Identify sentence boundaries for better chunking
            sentences = [sent.text for sent in doc.sents]
            text = ' '.join(sentences)
        
        return text
    
    def _create_chunks(self, text: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """
        Split text into chunks optimized for German language.
        
        Args:
            text: Cleaned document text
            metadata: Document metadata
            
        Returns:
            List of document chunks with metadata
        """
        if not text:
            return []
        
        # Split text into chunks
        texts = self.text_splitter.split_text(text)
        
        # Create document chunks with metadata
        chunks = []
        for i, chunk_text in enumerate(texts):
            chunk = DocumentChunk(
                text=chunk_text,
                metadata=metadata,
                chunk_id=i
            )
            chunks.append(chunk)
        
        return chunks